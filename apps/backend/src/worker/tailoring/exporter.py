import logging
import uuid
import os
from docx import Document
from fpdf import FPDF

logger = logging.getLogger(__name__)

class ResumeExporter:
    @staticmethod
    def export_all(tailored_resume: dict, variant_id: int):
        """
        Converts the finalized structural dict into .docx and .pdf
        """
        logger.info("Exporting DOCX and PDF...")
        
        base_name = f"tailored_resume_{variant_id}_{uuid.uuid4().hex[:8]}"
        docx_key = f"{base_name}.docx"
        pdf_key = f"{base_name}.pdf"
        
        docx_path = f"/tmp/{docx_key}"
        pdf_path = f"/tmp/{pdf_key}"
        
        # Export DOCX
        doc = Document()
        
        contact = tailored_resume.get("contact", {})
        doc.add_heading(contact.get("name", "Applicant"), 0)
        doc.add_paragraph(f"{contact.get('email', '')} | {contact.get('phone', '')} | {contact.get('location', '')}")
        
        summary = tailored_resume.get("summary", {})
        if summary:
            doc.add_heading(summary.get("title", ""), level=1)
            doc.add_paragraph(summary.get("summary_text", ""))
            
        skills = tailored_resume.get("skills", [])
        if skills:
            doc.add_heading("Technical Skills", level=1)
            doc.add_paragraph(", ".join(skills))
            
        for exp in tailored_resume.get("experience", []):
            doc.add_heading(f"{exp.get('title', '')} at {exp.get('company', '')}", level=2)
            doc.add_paragraph(f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}")
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style='List Bullet')
                
        doc.save(docx_path)
        
        # Export PDF (Lightweight ATS-Friendly)
        # Using FPDF which allows extremely deterministic minimal string pushing
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        pdf.set_font("Arial", "B", 16)
        pdf.cell(200, 10, txt=contact.get("name", "Applicant"), ln=True, align='C')
        pdf.set_font("Arial", size=10)
        pdf.cell(200, 5, txt=f"{contact.get('email', '')} | {contact.get('phone', '')} | {contact.get('location', '')}", ln=True, align='C')
        pdf.ln(5)
        
        if summary.get("title"):
            pdf.set_font("Arial", "B", 12)
            pdf.cell(200, 8, txt=summary.get("title", ""), ln=True)
            pdf.set_font("Arial", size=10)
            pdf.multi_cell(0, 5, txt=summary.get("summary_text", ""))
            pdf.ln(5)
            
        if skills:
            pdf.set_font("Arial", "B", 12)
            pdf.cell(200, 8, txt="Technical Skills", ln=True)
            pdf.set_font("Arial", size=10)
            pdf.multi_cell(0, 5, txt=", ".join(skills))
            pdf.ln(5)
            
        for exp in tailored_resume.get("experience", []):
            pdf.set_font("Arial", "B", 11)
            pdf.cell(200, 6, txt=f"{exp.get('title', '')} at {exp.get('company', '')}", ln=True)
            pdf.set_font("Arial", "I", 10)
            pdf.cell(200, 5, txt=f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}", ln=True)
            pdf.set_font("Arial", size=10)
            for bullet in exp.get("bullets", []):
                # Ensure ascii matching to avoid FPDF character errors naturally
                clean_bullet = bullet.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 5, txt=f"- {clean_bullet}")
            pdf.ln(3)

        pdf.output(pdf_path)
        
        return docx_key, pdf_key
